import os
import math
import re
import tempfile
import threading
from pathlib import Path
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from starlette.background import BackgroundTask
from faster_whisper import WhisperModel
from pydub import AudioSegment
from docx import Document

app = FastAPI()

# 精度重視。重い場合は環境変数 WHISPER_MODEL=medium で軽量化できます。
MODEL_SIZE = os.getenv("WHISPER_MODEL", "large-v3")
print(f"Loading Whisper model ({MODEL_SIZE})...- 精度重視モード")
model = WhisperModel(MODEL_SIZE, device="cpu", compute_type="int8")
print("Model loaded successfully!")

DOMAIN_PROMPT = (
    "これは日本語の電話応対の会話です。発言を省略せず、自然な句読点を付けてください。"
    "固有名詞と専門用語: まとめて光、So-net 光、ソネット光、オペレーター、"
    "受付担当、お客さま、契約者、回線、転用、事業者変更、工事費、違約金。"
)

REPLACEMENTS = {
    "ソネット 光": "So-net 光",
    "ソネット光": "So-net 光",
    "まとめてひかり": "まとめて光",
}

def format_time(seconds: float) -> str:
    m = math.floor(seconds / 60)
    s = math.floor(seconds % 60)
    return f"{m:02d}:{s:02d}"

def normalize_audio(audio: AudioSegment) -> AudioSegment:
    """Whisper向けに16kHz/monoへ統一し、小さすぎる音声だけ増幅する。"""
    audio = audio.set_frame_rate(16000).set_sample_width(2).set_channels(1)
    if audio.dBFS != float("-inf") and audio.dBFS < -24:
        audio = audio.apply_gain(min(12.0, -20.0 - audio.dBFS))
    return audio

def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    for wrong, correct in REPLACEMENTS.items():
        text = text.replace(wrong, correct)
    return text

def transcribe_channel(path: str, speaker: str):
    segments, _ = model.transcribe(
        path,
        language="ja",
        beam_size=8,
        best_of=8,
        patience=1.2,
        temperature=0.0,
        initial_prompt=DOMAIN_PROMPT,
        condition_on_previous_text=True,
        vad_filter=True,
        vad_parameters={"min_silence_duration_ms": 500, "speech_pad_ms": 250},
        no_speech_threshold=0.6,
        log_prob_threshold=-1.0,
        compression_ratio_threshold=2.4,
    )
    result = []
    last_text = ""
    for seg in segments:
        text = clean_text(seg.text)
        # 無音区間で起きやすい同一文の連続出力を除く。
        if text and text != last_text:
            result.append({"start": seg.start, "speaker": speaker, "text": text})
            last_text = text
    return result

@app.get("/", response_class=HTMLResponse)
async def index():
    return """
    <!DOCTYPE html>
    <html lang="ja">
    <head>
        <meta charset="UTF-8">
        <title>ローカル高精度音声文字起こしシステム</title>
        <style>
            :root {
                --bg-color: #f1f5f9;
                --card-bg: #ffffff;
                --primary-color: #2563eb;
                --primary-hover: #1d4ed8;
                --text-color: #1e293b;
                --border-color: #cbd5e1;
            }
            body {
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
                background-color: var(--bg-color);
                color: var(--text-color);
                margin: 0;
                padding: 15px;
                display: flex;
                justify-content: center;
                align-items: center;
                min-height: 100vh;
                box-sizing: border-box;
            }
            .container {
                width: 100%;
                max-width: 800px;
                background: var(--card-bg);
                padding: 25px;
                border-radius: 12px;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
                border: 1px solid #e2e8f0;
            }
            .header-flex {
                border-bottom: 2px solid #f1f5f9;
                padding-bottom: 12px;
                margin-bottom: 20px;
            }
            h2 {
                margin: 0;
                font-size: 1.4rem;
                color: #0f172a;
            }
            .subtitle {
                font-size: 0.85rem;
                color: #64748b;
                margin-top: 4px;
            }
            .drop-zone {
                border: 2px dashed var(--border-color);
                border-radius: 8px;
                padding: 22px;
                text-align: center;
                background: #f8fafc;
                cursor: pointer;
                transition: all 0.2s;
                margin-bottom: 15px;
            }
            .drop-zone.dragover {
                background: #eff6ff;
                border-color: var(--primary-color);
            }
            .drop-zone p {
                margin: 0 0 6px;
                color: #475569;
                font-weight: 500;
                font-size: 0.95rem;
            }
            .file-info {
                font-weight: bold;
                color: var(--primary-color);
                margin-top: 4px;
                font-size: 0.9rem;
            }
            .btn-container {
                display: flex;
                gap: 10px;
                margin-bottom: 15px;
            }
            button {
                background-color: var(--primary-color);
                color: white;
                border: none;
                padding: 10px 16px;
                font-size: 0.95rem;
                font-weight: 600;
                border-radius: 6px;
                cursor: pointer;
                transition: background-color 0.2s;
                flex: 1;
            }
            button:hover {
                background-color: var(--primary-hover);
            }
            button.secondary {
                background-color: #475569;
            }
            button.secondary:hover {
                background-color: #334155;
            }
            button.success {
                background-color: #0d9488;
            }
            button.success:hover {
                background-color: #0f766e;
            }
            #status {
                font-weight: 600;
                color: #d97706;
                margin-bottom: 8px;
                font-size: 0.9rem;
                min-height: 20px;
            }
            textarea {
                width: 100%;
                height: 380px;
                padding: 12px;
                border: 1px solid var(--border-color);
                border-radius: 6px;
                font-family: inherit;
                font-size: 0.95rem;
                line-height: 1.5;
                resize: vertical;
                box-sizing: border-box;
                background: #fdfdfd;
                color: #1e293b;
            }
            input[type="file"] {
                display: none;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header-flex">
                <h2>ローカル高精度音声文字起こしシステム</h2>
                <div class="subtitle">左右チャンネル分離 ＆ 専門用語補正モード</div>
            </div>
            
            <div class="drop-zone" id="dropZone" onclick="document.getElementById('audioFile').click()">
                <p>音声ファイルをここにドラッグ＆ドロップしてください<br><span style="font-size: 0.8rem; color: #64748b;">AP音声→客音声の順に2ファイル選択（1ファイルでも可）</span></p>
                <div class="file-info" id="fileInfo">ファイルが選択されていません</div>
                <input type="file" id="audioFile" accept="audio/*,.mp3,.wav,.m4a,.aac,.flac,.ogg" multiple>
            </div>

            <div class="btn-container">
                <button onclick="transcribe()">文字起こし開始</button>
                <button class="secondary" onclick="copyResult()">結果をコピー</button>
                <button class="success" onclick="downloadWord()">Wordで保存</button>
            </div>

            <div id="status"></div>
            <textarea id="result" placeholder="ここに文字起こし結果が表示されます..."></textarea>
        </div>

        <script>
            const dropZone = document.getElementById('dropZone');
            const fileInput = document.getElementById('audioFile');
            const fileInfo = document.getElementById('fileInfo');
            const status = document.getElementById('status');
            const result = document.getElementById('result');
            let selectedFiles = [];

            ['dragenter', 'dragover'].forEach(eventName => {
                dropZone.addEventListener(eventName, (e) => {
                    e.preventDefault();
                    dropZone.classList.add('dragover');
                }, false);
            });

            ['dragleave', 'drop'].forEach(eventName => {
                dropZone.addEventListener(eventName, (e) => {
                    e.preventDefault();
                    dropZone.classList.remove('dragover');
                }, false);
            });

            dropZone.addEventListener('drop', (e) => {
                const files = e.dataTransfer.files;
                if (files.length > 0) {
                    fileInput.files = files;
                    handleFileSelection(files);
                }
            }, false);

            fileInput.addEventListener('change', () => {
                if (fileInput.files.length > 0) {
                    handleFileSelection(fileInput.files);
                }
            });

            function handleFileSelection(files) {
                selectedFiles = Array.from(files).slice(0, 2);
                if (selectedFiles.length === 2) {
                    fileInfo.innerText = `AP: ${selectedFiles[0].name} / 客: ${selectedFiles[1].name}`;
                } else {
                    fileInfo.innerText = `選択中: ${selectedFiles[0].name}（左右チャンネルを自動判定）`;
                }
            }

            async function transcribe() {
                if (selectedFiles.length === 0) {
                    alert('音声ファイルを選択してください');
                    return;
                }

                const formData = new FormData();
                selectedFiles.forEach(file => formData.append('files', file));

                status.innerText = '高精度解析中（チャンネル分離・専門用語補正）...';
                result.value = '';

                try {
                    const response = await fetch('/transcribe', {
                        method: 'POST',
                        body: formData
                    });
                    const data = await response.json();
                    if (response.ok) {
                        result.value = data.text;
                        status.innerText = '文字起こしが完了しました';
                    } else {
                        status.innerText = 'エラーが発生しました';
                        result.value = data.detail;
                    }
                } catch (err) {
                    status.innerText = '通信エラーが発生しました';
                    result.value = err;
                }
            }

            function copyResult() {
                if (!result.value) {
                    alert('コピーするテキストがありません');
                    return;
                }
                navigator.clipboard.writeText(result.value).then(() => {
                    alert('クリップボードにコピーしました');
                });
            }

            async function downloadWord() {
                const text = result.value;
                if (!text) {
                    alert('保存するテキストがありません');
                    return;
                }

                status.innerText = 'Wordファイルを作成中...';

                try {
                    const response = await fetch('/download-word', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ text: text })
                    });

                    if (response.ok) {
                        const blob = await response.blob();
                        const url = window.URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.href = url;
                        a.download = '高精度文字起こし結果.docx';
                        document.body.appendChild(a);
                        a.click();
                        a.remove();
                        status.innerText = 'Wordファイルのダウンロードが完了しました';
                    } else {
                        status.innerText = 'Wordの作成に失敗しました';
                    }
                } catch (err) {
                    status.innerText = 'エラーが発生しました';
                }
            }
            window.addEventListener('beforeunload', () => {
                navigator.sendBeacon('/shutdown');
            });
        </script>
    </body>
    </html>
    """

@app.post("/transcribe")
async def transcribe_audio(files: list[UploadFile] = File(...)):
    try:
        if not 1 <= len(files) <= 2:
            raise HTTPException(status_code=400, detail="音声は1つまたは2つ選択してください。")
        with tempfile.TemporaryDirectory(prefix="transcribe_") as temp_dir:
            all_segments = []
            sources = []
            for index, upload in enumerate(files):
                suffix = Path(upload.filename or "audio").suffix
                source_path = os.path.join(temp_dir, f"source_{index}{suffix}")
                with open(source_path, "wb") as f:
                    while chunk := await upload.read(1024 * 1024):
                        f.write(chunk)
                sources.append(AudioSegment.from_file(source_path))

            if len(sources) == 2:
                for audio, filename, speaker in (
                    (sources[0], "ap.wav", "AP"),
                    (sources[1], "customer.wav", "客"),
                ):
                    path = os.path.join(temp_dir, filename)
                    normalize_audio(audio).export(path, format="wav")
                    all_segments.extend(transcribe_channel(path, speaker))
            else:
                channels = sources[0].split_to_mono()
                if len(channels) >= 2:
                    for audio, filename, speaker in (
                        (channels[0], "ap.wav", "AP"),
                        (channels[1], "customer.wav", "客"),
                    ):
                        path = os.path.join(temp_dir, filename)
                        normalize_audio(audio).export(path, format="wav")
                        all_segments.extend(transcribe_channel(path, speaker))
                else:
                    mono_path = os.path.join(temp_dir, "mono.wav")
                    normalize_audio(sources[0]).export(mono_path, format="wav")
                    all_segments.extend(transcribe_channel(mono_path, "話者"))

        # タイムスタンプ順に時系列ソート
        all_segments.sort(key=lambda x: x["start"])

        safe_names = " / ".join(Path(f.filename or "音声ファイル").name for f in files)
        full_text = f"通話文字起こし\n対象ファイル: {safe_names}\n\n【文字起こしログ】\n\n"
        for seg in all_segments:
            time_tag = format_time(seg["start"])
            full_text += f"[{seg['speaker']} {time_tag}] {seg['text']}\n"

        return {"text": full_text}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/download-word")
async def download_word(data: dict):
    text = data.get("text", "")
    doc = Document()
    doc.add_heading('通話文字起こし', 0)
    for line in text.splitlines():
        doc.add_paragraph(line)

    fd, file_path = tempfile.mkstemp(prefix="transcript_", suffix=".docx")
    os.close(fd)
    doc.save(file_path)

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="高精度文字起こし結果.docx",
        background=BackgroundTask(lambda: os.path.exists(file_path) and os.remove(file_path)),
    )

@app.post("/shutdown")
async def shutdown():
    threading.Timer(1.0, lambda: os._exit(0)).start()
    return {"status": "closing"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)

